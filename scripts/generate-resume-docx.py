"""Generate Kunal-Varshney-Resume.docx from current résumé content."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parent.parent / "assets" / "Kunal-Varshney-Resume.docx"

FONT = "Calibri"
BODY = 11
SMALL = 10.5
HEAD = 11.5
NAME = 22


def set_margins(doc: Document) -> None:
    for sec in doc.sections:
        sec.top_margin = Inches(0.55)
        sec.bottom_margin = Inches(0.55)
        sec.left_margin = Inches(0.55)
        sec.right_margin = Inches(0.55)


def style_paragraph(p, size=BODY, bold=False, italic=False, space_after=Pt(4)):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = space_after
    for run in p.runs:
        run.font.name = FONT
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(HEAD)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.name = FONT
            run.font.size = Pt(BODY)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_job(doc: Document, title: str, dates: str, org: str, bullets: list[str]) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.name = FONT
    r1.font.size = Pt(BODY)
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r2 = p.add_run(f"\t{dates}")
    r2.bold = True
    r2.font.name = FONT
    r2.font.size = Pt(SMALL)
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    p.paragraph_format.space_after = Pt(0)

    o = doc.add_paragraph()
    ro = o.add_run(org)
    ro.italic = True
    ro.font.name = FONT
    ro.font.size = Pt(SMALL)
    ro._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    o.paragraph_format.space_after = Pt(2)

    add_bullets(doc, bullets)


def main() -> None:
    doc = Document()
    set_margins(doc)

    # Name block
    p = doc.add_paragraph()
    r = p.add_run("Kunal Varshney")
    r.bold = True
    r.font.size = Pt(NAME)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph("Web Designer · AI-Assisted Development")
    style_paragraph(p, size=BODY, bold=True, space_after=Pt(2))

    p = doc.add_paragraph(
        "kunalvrshn@gmail.com | +91 70176 62533 | India (Remote)\n"
        "lilkunal.github.io | linkedin.com/in/kunalvrshn | github.com/lilkunal"
    )
    style_paragraph(p, size=SMALL, space_after=Pt(10))

    # Skills
    add_heading(doc, "Technical Skills")
    skills = [
        ("Web Design", "HTML, CSS, JavaScript, responsive layout, landing pages, SEO, Photoshop, Git, GitHub Pages"),
        ("AI Workflow", "Claude, Cursor — faster build with personal review before launch"),
        ("Marketing", "Digital marketing, catalogues, trade pricing, Excel, client communication"),
        ("Support", "PBX, remote troubleshooting, ticketing, CRM, Oracle WebLogic, Jenkins, log analysis"),
        ("Soft Skills", "Communication, time management, presentation, problem-solving, decision-making"),
    ]
    for label, val in skills:
        p = doc.add_paragraph()
        rl = p.add_run(f"{label}: ")
        rl.bold = True
        rl.font.name = FONT
        rl.font.size = Pt(BODY)
        rl._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        rv = p.add_run(val)
        rv.font.name = FONT
        rv.font.size = Pt(BODY)
        rv._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15

    # Experience page 1
    add_heading(doc, "Professional Experience")
    add_job(
        doc,
        "Freelance Web Designer",
        "Jan 2025 – Present",
        "Independent · Remote · JAI Home Care & small-business clients",
        [
            "Design and build marketing websites and landing pages from brief through launch.",
            "Use Claude and Cursor for layout and iteration; verify code, copy, and mobile behaviour before go-live.",
            "Delivered JAI Home Care landing page; scope, pricing, and post-launch support agreed in writing.",
        ],
    )
    add_job(
        doc,
        "Marketing & Online Specialist",
        "Jan 2025 – Present",
        "Padma Enterprises (Padma Lights) · India",
        [
            "Lead website, digital marketing, and online presence for cast-aluminium outdoor lighting.",
            "Manage 200+ model catalogue, trade pricing, and international orders in eight currencies.",
            "Coordinate product photography and catalogue updates with manufacturing and export teams.",
            "Grew the business online for nine years before assuming this role full-time in 2025.",
        ],
    )
    add_job(
        doc,
        "Technical Support Advisor",
        "Aug 2024 – Dec 2024",
        "Teleperformance",
        [
            "Diagnosed and resolved technical issues via remote troubleshooting and ticketing in PBX environments.",
            "Maintained service quality in a high-volume setting with changing tools and procedures.",
            "Documented resolutions and follow-ups in CRM for handoffs and audit trails.",
        ],
    )

    doc.add_page_break()

    p = doc.add_paragraph()
    r = p.add_run("Kunal Varshney")
    r.bold = True
    r.font.size = Pt(HEAD)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    p.paragraph_format.space_after = Pt(8)

    add_heading(doc, "Professional Experience")
    add_job(
        doc,
        "Customer Support Executive",
        "Oct 2023 – Jan 2024",
        "Webhelp Pvt. Ltd.",
        [
            "Customer service via chat and email — product enquiries, complaints, and troubleshooting.",
            "Maintained CRM records; escalated complex cases and coordinated with teams until resolution.",
        ],
    )
    add_job(
        doc,
        "Content Moderator",
        "Oct 2021 – Apr 2023",
        "Concentrix Daksh Pvt. Ltd.",
        [
            "Reviewed content for Google platforms including YouTube Kids using internal moderation tools.",
            "Prepared team reports; created transcripts and aligned content for production workflows.",
        ],
    )
    add_job(
        doc,
        "Technical Support Trainee",
        "Feb 2021 – Jun 2021",
        "Flex Pvt. Ltd.",
        ["Supported Oracle WebLogic 11g and Jenkins monitoring; analysed logs and produced health-check reports."],
    )
    add_job(
        doc,
        "Intern",
        "Jun 2019 – Aug 2019",
        "Parag Dairy · UP Rajya Vidyut Utpadan Nigam Ltd.",
        [
            "Quality-control operations and delivery records at Parag Dairy; progress reports at thermal power station."
        ],
    )

    add_heading(doc, "Selected Projects")
    projects = [
        ("Padma Enterprises — Padma Lights (padmalights.com)", "Marketing & online operations · Ongoing",
         "Website and digital ops — 200+ SKU catalogue, trade pricing, export content."),
        ("JAI Home Care Services · 2025", "",
         "Home-care landing page — services, contact, phone-optimised layout for local enquiries."),
        ("Personal Portfolio (lilkunal.github.io) · 2025", "",
         "Professional portfolio — AI-assisted build with full personal QA on design and content."),
    ]
    for title, meta, desc in projects:
        p = doc.add_paragraph()
        rt = p.add_run(title)
        rt.bold = True
        rt.font.name = FONT
        rt.font.size = Pt(BODY)
        rt._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        if meta:
            rm = p.add_run(f" — {meta}")
            rm.italic = True
            rm.font.name = FONT
            rm.font.size = Pt(SMALL)
            rm._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        p.paragraph_format.space_after = Pt(1)
        d = doc.add_paragraph(desc)
        style_paragraph(d, space_after=Pt(4))

    add_heading(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run("B.Tech in Computer Science\n")
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(BODY)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r2 = p.add_run("Aligarh College of Engineering & Technology · Graduated 2020")
    r2.font.name = FONT
    r2.font.size = Pt(BODY)
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    p.paragraph_format.space_after = Pt(6)

    add_heading(doc, "Certifications")
    add_bullets(
        doc,
        [
            "Google IT Support Professional",
            "Google Digital Garage — Digital Marketing",
            "Excel Data Analysis (Coursera)",
            "ChatGPT for Beginners (Great Learning)",
        ],
    )

    add_heading(doc, "Honours & Activities")
    add_bullets(
        doc,
        [
            "1st Prize — CS Quiz, Vivekananda College of Technology & Management",
            "2nd Prize — Debate Competition, Varshney College",
            "Workshops: Networking (Aptron), Blockchain & ML (Training Basket)",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
